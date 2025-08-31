#!/usr/bin/env python3
"""
File Upload and Processing System for HRM-Gemini AI
Handles document and image uploads with intelligent processing
"""

import sys
import os
import base64
import mimetypes
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import tempfile
import shutil

# Add project root to path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

class FileUploadSystem:
    """Handles file uploads and processing for the AI system"""

    def __init__(self, upload_dir: str = "uploads", memory_system=None):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        self.memory_system = memory_system

        # File type handlers
        self.file_handlers = {
            'text': self._process_text_file,
            'document': self._process_document_file,
            'image': self._process_image_file,
            'code': self._process_code_file,
            'data': self._process_data_file
        }

        # Supported file types
        self.supported_types = {
            # Text files
            '.txt': 'text', '.md': 'text', '.rtf': 'text',

            # Documents
            '.pdf': 'document', '.docx': 'document', '.doc': 'document',
            '.xlsx': 'data', '.xls': 'data', '.csv': 'data',

            # Code files
            '.py': 'code', '.js': 'code', '.java': 'code', '.cpp': 'code',
            '.c': 'code', '.html': 'text', '.css': 'text', '.json': 'data',

            # Images
            '.jpg': 'image', '.jpeg': 'image', '.png': 'image',
            '.gif': 'image', '.bmp': 'image', '.tiff': 'image', '.webp': 'image'
        }

    def upload_file(self, file_path: str, description: str = "") -> Dict[str, Any]:
        """Upload and process a file"""
        try:
            file_path_obj = Path(file_path)

            if not file_path_obj.exists():
                return {
                    'success': False,
                    'error': f'File not found: {file_path}',
                    'file_info': None
                }

            # Validate file type
            file_type = self._get_file_type(file_path)
            if not file_type:
                return {
                    'success': False,
                    'error': f'Unsupported file type: {file_path_obj.suffix}',
                    'file_info': None
                }

            # Copy file to upload directory
            upload_path = self._store_file(file_path_obj)

            # Process file content
            processed_content = self._process_file(upload_path, file_type)

            # Store in memory system
            if self.memory_system:
                memory_result = self.memory_system.process_file_upload(
                    str(upload_path),
                    file_type
                )

                # Add additional metadata
                metadata = {
                    'original_path': str(file_path_obj),
                    'upload_path': str(upload_path),
                    'file_size': file_path_obj.stat().st_size,
                    'description': description,
                    'processed_content_length': len(processed_content)
                }

                self.memory_system.store_memory(
                    f"File uploaded: {file_path_obj.name} - {description}",
                    memory_type="file_upload",
                    metadata=metadata,
                    importance=0.8
                )

            file_info = {
                'name': file_path_obj.name,
                'size': file_path_obj.stat().st_size,
                'type': file_type,
                'upload_path': str(upload_path),
                'processed_content_length': len(processed_content)
            }

            return {
                'success': True,
                'message': f'Successfully uploaded and processed: {file_path_obj.name}',
                'file_info': file_info,
                'content_preview': processed_content[:500] + "..." if len(processed_content) > 500 else processed_content
            }

        except Exception as e:
            return {
                'success': False,
                'error': f'Upload failed: {e}',
                'file_info': None
            }

    def _get_file_type(self, file_path: str) -> Optional[str]:
        """Determine file type based on extension"""
        extension = Path(file_path).suffix.lower()
        return self.supported_types.get(extension)

    def _store_file(self, file_path: Path) -> Path:
        """Store uploaded file in upload directory"""
        import uuid

        # Generate unique filename
        unique_name = f"{uuid.uuid4()}_{file_path.name}"
        upload_path = self.upload_dir / unique_name

        # Copy file
        shutil.copy2(file_path, upload_path)

        return upload_path

    def _process_file(self, file_path: Path, file_type: str) -> str:
        """Process file based on its type"""
        if file_type in self.file_handlers:
            return self.file_handlers[file_type](file_path)
        else:
            return f"[Unsupported file type: {file_type}]"

    def _process_text_file(self, file_path: Path) -> str:
        """Process text files"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            # Basic text cleaning
            content = content.strip()
            content = '\n'.join(line for line in content.split('\n') if line.strip())

            return content

        except Exception as e:
            return f"[Text processing error: {e}]"

    def _process_document_file(self, file_path: Path) -> str:
        """Process document files (PDF, DOCX, etc.)"""
        extension = file_path.suffix.lower()

        if extension == '.pdf':
            return self._process_pdf_file(file_path)
        elif extension in ['.docx', '.doc']:
            return self._process_docx_file(file_path)
        else:
            return f"[Unsupported document type: {extension}]"

    def _process_pdf_file(self, file_path: Path) -> str:
        """Process PDF files"""
        try:
            import PyPDF2

            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"

                return text.strip()

        except ImportError:
            return "[PDF processing requires PyPDF2. Install with: pip install PyPDF2]"
        except Exception as e:
            return f"[PDF processing error: {e}]"

    def _process_docx_file(self, file_path: Path) -> str:
        """Process DOCX files"""
        try:
            from docx import Document

            doc = Document(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"

            return text.strip()

        except ImportError:
            return "[DOCX processing requires python-docx. Install with: pip install python-docx]"
        except Exception as e:
            return f"[DOCX processing error: {e}]"

    def _process_image_file(self, file_path: Path) -> str:
        """Process image files"""
        try:
            # For images, we'll extract basic metadata and prepare for AI processing
            import PIL.Image
            import io

            img = PIL.Image.open(file_path)

            # Get basic image info
            width, height = img.size
            format_type = img.format
            mode = img.mode

            # Convert to base64 for storage/transmission
            buffer = io.BytesIO()
            img.save(buffer, format=format_type or 'PNG')
            img_base64 = base64.b64encode(buffer.getvalue()).decode()

            metadata = f"""
Image Information:
- Filename: {file_path.name}
- Dimensions: {width}x{height}
- Format: {format_type}
- Color Mode: {mode}
- File Size: {file_path.stat().st_size} bytes

[Image data encoded in base64 - ready for AI processing]
Base64 Preview: {img_base64[:100]}...
"""

            return metadata.strip()

        except ImportError:
            return "[Image processing requires PIL/Pillow. Install with: pip install Pillow]"
        except Exception as e:
            return f"[Image processing error: {e}]"

    def _process_code_file(self, file_path: Path) -> str:
        """Process code files with syntax highlighting info"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            # Add language identification
            extension = file_path.suffix.lower()
            lang_map = {
                '.py': 'Python',
                '.js': 'JavaScript',
                '.java': 'Java',
                '.cpp': 'C++',
                '.c': 'C',
                '.html': 'HTML',
                '.css': 'CSS'
            }

            language = lang_map.get(extension, 'Unknown')
            header = f"Code File: {file_path.name}\nLanguage: {language}\n{'='*50}\n\n"

            return header + content

        except Exception as e:
            return f"[Code processing error: {e}]"

    def _process_data_file(self, file_path: Path) -> str:
        """Process data files (CSV, JSON, Excel)"""
        extension = file_path.suffix.lower()

        if extension == '.json':
            return self._process_json_file(file_path)
        elif extension == '.csv':
            return self._process_csv_file(file_path)
        elif extension in ['.xlsx', '.xls']:
            return self._process_excel_file(file_path)
        else:
            return f"[Unsupported data file type: {extension}]"

    def _process_json_file(self, file_path: Path) -> str:
        """Process JSON files"""
        try:
            import json
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            # Convert to readable format
            content = f"JSON File: {file_path.name}\n{'='*50}\n\n"
            content += json.dumps(data, indent=2)

            return content

        except Exception as e:
            return f"[JSON processing error: {e}]"

    def _process_csv_file(self, file_path: Path) -> str:
        """Process CSV files"""
        try:
            import csv
            content = f"CSV File: {file_path.name}\n{'='*50}\n\n"

            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                rows = list(reader)

                if rows:
                    # Show header
                    content += f"Headers: {', '.join(rows[0])}\n"
                    content += f"Total rows: {len(rows)}\n\n"

                    # Show first few rows as preview
                    preview_rows = min(5, len(rows))
                    for i in range(preview_rows):
                        content += f"Row {i+1}: {', '.join(rows[i])}\n"

                    if len(rows) > preview_rows:
                        content += f"\n... and {len(rows) - preview_rows} more rows"

            return content

        except Exception as e:
            return f"[CSV processing error: {e}]"

    def _process_excel_file(self, file_path: Path) -> str:
        """Process Excel files"""
        try:
            import pandas as pd

            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None)

            content = f"Excel File: {file_path.name}\n{'='*50}\n\n"
            content += f"Sheets found: {', '.join(excel_data.keys())}\n\n"

            for sheet_name, df in excel_data.items():
                content += f"Sheet: {sheet_name}\n"
                content += f"Shape: {df.shape[0]} rows × {df.shape[1]} columns\n"

                if not df.empty:
                    content += f"Columns: {', '.join(df.columns.tolist())}\n"
                    content += f"Preview:\n{df.head(3).to_string()}\n\n"

            return content

        except ImportError:
            return "[Excel processing requires pandas and openpyxl. Install with: pip install pandas openpyxl]"
        except Exception as e:
            return f"[Excel processing error: {e}]"

    def list_uploaded_files(self) -> List[Dict[str, Any]]:
        """List all uploaded files"""
        files_info = []

        for file_path in self.upload_dir.glob('*'):
            if file_path.is_file():
                stat = file_path.stat()
                files_info.append({
                    'name': file_path.name,
                    'path': str(file_path),
                    'size': stat.st_size,
                    'modified': stat.st_mtime,
                    'type': self._get_file_type(str(file_path)) or 'unknown'
                })

        # Sort by modification time (newest first)
        files_info.sort(key=lambda x: x['modified'], reverse=True)

        return files_info

    def search_files(self, query: str) -> List[Dict[str, Any]]:
        """Search uploaded files by name or content"""
        query_lower = query.lower()
        matching_files = []

        for file_info in self.list_uploaded_files():
            if query_lower in file_info['name'].lower():
                matching_files.append(file_info)

        return matching_files

    def get_file_content(self, filename: str) -> Optional[str]:
        """Get processed content of a specific file"""
        file_path = self.upload_dir / filename

        if file_path.exists():
            file_type = self._get_file_type(str(file_path))
            return self._process_file(file_path, file_type)

        return None

    def cleanup_old_files(self, days_old: int = 30) -> int:
        """Clean up files older than specified days"""
        import time

        cutoff_time = time.time() - (days_old * 24 * 60 * 60)
        deleted_count = 0

        for file_path in self.upload_dir.glob('*'):
            if file_path.is_file() and file_path.stat().st_mtime < cutoff_time:
                file_path.unlink()
                deleted_count += 1

        return deleted_count
